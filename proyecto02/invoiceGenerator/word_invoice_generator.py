"""
Generador de facturas usando plantilla de Word - versión mejorada
Ahora crea la tabla desde cero si no existe
"""
from docx import Document
from datetime import datetime
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import os

path_template = "invoiceGenerator\\FACTURA_BOCETO.docx"

class WordInvoiceGenerator:
    def __init__(self, template_path=path_template):
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"No se encontró la plantilla: {template_path}")
        self.template_path = template_path

    def replace_all_placeholders(self, doc, replacements):
        """Reemplaza texto en párrafos, tablas y cuadros de texto."""
        # Reemplazar en párrafos normales
        for p in doc.paragraphs:
            for marcador, valor in replacements.items():
                if marcador in p.text:
                    for run in p.runs:
                        if marcador in run.text:
                            run.text = run.text.replace(marcador, valor)

        # Reemplazar dentro de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for marcador, valor in replacements.items():
                            if marcador in p.text:
                                for run in p.runs:
                                    if marcador in run.text:
                                        run.text = run.text.replace(marcador, valor)

        # Reemplazar dentro de cuadros de texto (text boxes)
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for txbx in doc.element.findall('.//w:txbxContent', nsmap):
            for t in txbx.findall('.//w:t', nsmap):
                if t.text:
                    for marcador, valor in replacements.items():
                        if marcador in t.text:
                            t.text = t.text.replace(marcador, valor)

    def calculate_totals(self, items):
        subtotal = sum(item['quantity'] * item['unit_price'] for item in items)
        descuento = 0
        iva = subtotal * 0.16
        total = subtotal - descuento + iva
        return {'subtotal': subtotal, 'descuento': descuento, 'iva': iva, 'total': total}

    def find_table_marker(self, doc):
        """Busca el párrafo donde debe insertarse la tabla de items."""
        for i, p in enumerate(doc.paragraphs):
            # Buscar el texto que indica donde va la tabla
            if 'Cantidad' in p.text and 'Descripción' in p.text:
                return i
        return None

    def create_items_table(self, doc, items, insert_position=None):
        """Crea una tabla formateada para los items de factura."""
        
        # Si no se especifica posición, buscar automáticamente
        if insert_position is None:
            insert_position = self.find_table_marker(doc)
        
        # Crear la tabla: 1 fila de encabezado + filas de items
        table = doc.add_table(rows=1 + len(items), cols=5)
        table.style = 'Table Grid'
        
        # Configurar encabezados
        headers = ['Cantidad', 'U. de medida', 'Descripción', 'Precio unitario', 'Importe']
        header_cells = table.rows[0].cells
        
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            # Formato de encabezado
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Color de fondo (gris claro)
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Llenar filas con items
        for idx, item in enumerate(items, start=1):
            row = table.rows[idx]
            cells = row.cells
            
            # Cantidad
            cells[0].text = f"{item['quantity']:.2f}"
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Unidad
            cells[1].text = item['unit']
            cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Descripción
            cells[2].text = item['description']
            
            # Precio unitario
            cells[3].text = f"${item['unit_price']:,.2f}"
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Importe
            importe = item['quantity'] * item['unit_price']
            cells[4].text = f"${importe:,.2f}"
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Aplicar formato a todas las celdas
            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        
        # Ajustar anchos de columnas
        widths = [Inches(0.8), Inches(1.0), Inches(3.0), Inches(1.2), Inches(1.2)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        
        return table

    def remove_placeholder_table_text(self, doc):
        """Elimina el texto de la tabla placeholder en markdown."""
        to_remove = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            # Identificar líneas de la tabla markdown
            if (text.startswith('---') or 
                text.startswith('**Cantidad**') or
                text.startswith('**U. de medida**') or
                'Descripción' in text and '**' in text or
                len(text) > 0 and all(c in '-| ' for c in text)):
                to_remove.append(p)
        
        # Eliminar párrafos identificados
        for p in to_remove:
            p_element = p._element
            p_element.getparent().remove(p_element)

    def detect_placeholders(self, doc):
        """Detecta automáticamente todos los placeholders {{...}} en el documento."""
        placeholders = set()
        
        # Buscar en párrafos
        for p in doc.paragraphs:
            placeholders.update(re.findall(r"\{\{.*?\}\}", p.text))
        
        # Buscar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        placeholders.update(re.findall(r"\{\{.*?\}\}", p.text))
        
        # Buscar en cuadros de texto
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        for txbx in doc.element.findall('.//w:txbxContent', nsmap):
            for t in txbx.findall('.//w:t', nsmap):
                if t.text:
                    placeholders.update(re.findall(r"\{\{.*?\}\}", t.text))
        
        return placeholders

    def convert_to_pdf(self, docx_path, pdf_path):
        try:
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            print(f"✅ PDF generado: {pdf_path}")
            return pdf_path
        except Exception as e:
            print(f"⚠️ No se pudo convertir a PDF: {e}")
            return None

    def generate_invoice(self, client_data, items=None, output_docx="factura.docx", generate_pdf=True):
        """Genera la factura completa."""
        doc = Document(self.template_path)

        if items is None:
            items = [
                {'quantity': 1, 'unit': 'Servicio', 'description': 'Desarrollo de Software', 'unit_price': 15000.00}
            ]

        totals = self.calculate_totals(items)
        fecha_actual = datetime.now()
        folio_fiscal = fecha_actual.strftime("%Y%m%d%H%M%S")

        # Detectar placeholders existentes
        placeholders = self.detect_placeholders(doc)

        # Generar reemplazos automáticos
        replacements = {
            '{{NOMBRE}}': client_data.get('fiscal_name', ''),
            '{{RFC}}': client_data.get('rfc', ''),
            '{{DIRECCION}}': client_data.get('address', ''),
            '{{REGIMEN FISCAL}}': client_data.get('tax_regime', ''),
            '{{SUB TOTAL}}': f"${totals['subtotal']:,.2f}",
            '{{IVA}}': f"${totals['iva']:,.2f}",
            '{{TOTAL}}': f"${totals['total']:,.2f}",
            '{{FECHA}}': fecha_actual.strftime("%d/%m/%Y"),
            '{{HORA}}': fecha_actual.strftime("%H:%M:%S"),
            '{{FACTURA}}': folio_fiscal,
        }

        # Solo reemplaza placeholders que existan
        replacements = {k: v for k, v in replacements.items() if k in placeholders}

        # Reemplazar textos
        self.replace_all_placeholders(doc, replacements)

        # Eliminar texto placeholder de la tabla markdown
        self.remove_placeholder_table_text(doc)

        # Crear tabla de items (se agrega al final del documento)
        self.create_items_table(doc, items)

        # Guardar Word
        doc.save(output_docx)
        print(f"✅ Word generado: {output_docx}")

        # Generar PDF si se solicita
        pdf_path = None
        if generate_pdf:
            pdf_path = output_docx.replace('.docx', '.pdf')
            pdf_path = self.convert_to_pdf(output_docx, pdf_path)

        return {'docx': output_docx, 'pdf': pdf_path, 'totals': totals}


# Función auxiliar para shading
def nsdecls(*prefixes):
    """Genera namespace declarations para XML."""
    return ' '.join(f'xmlns:{p}="http://schemas.openxmlformats.org/wordprocessingml/2006/main"' 
                    if p == 'w' else '' for p in prefixes)

def parse_xml(xml_string):
    """Parse XML string."""
    from lxml import etree
    return etree.fromstring(xml_string)